from docx import Document
from docx.shared import Inches # Optional: If you need specific formatting like indentation
from io import BytesIO
import datetime # Keep for potential future use, though not directly used now

def create_docx(article_dict: dict):
    """
    Convert a dictionary conforming to the SaaS User Documentation schema
    into a DOCX document in memory.
    """
    doc = Document()

    # --- Document Structure based on JSON Schema ---

    # Title (string)
    title = article_dict.get('Title', '').strip() # Changed 'title' to 'Title'
    if title:
        doc.add_heading(title, level=0)

    # Subtitle (string)
    subtitle = article_dict.get('Subtitle', '').strip() # Changed 'subtitle' to 'Subtitle'
    if subtitle:
        doc.add_heading(subtitle, level=1) # Use level 1 for subtitle

    # Introduction (object)
    introduction = article_dict.get('Introduction') # Changed 'introduction' to 'Introduction'
    if introduction and isinstance(introduction, dict):
        doc.add_heading("Introduction", level=1) # Add a standard heading

        # Introduction Paragraphs (array of strings)
        intro_paragraphs = introduction.get('paragraphs')
        if intro_paragraphs and isinstance(intro_paragraphs, list):
            for para in intro_paragraphs:
                if isinstance(para, str) and para.strip():
                    doc.add_paragraph(para.strip())

        # Introduction Prerequisites (array of strings)
        prerequisites = introduction.get('prerequisites')
        if prerequisites and isinstance(prerequisites, list):
            doc.add_heading("Prerequisites", level=2) # Sub-heading
            for prereq in prerequisites:
                if isinstance(prereq, str) and prereq.strip():
                    doc.add_paragraph(prereq.strip(), style='List Bullet')

        # Introduction Outcomes (array of strings)
        outcomes = introduction.get('outcomes')
        if outcomes and isinstance(outcomes, list):
            doc.add_heading("Learning Outcomes", level=2) # Sub-heading
            for outcome in outcomes:
                if isinstance(outcome, str) and outcome.strip():
                    doc.add_paragraph(outcome.strip(), style='List Bullet')

    # Features (array of strings)
    features = article_dict.get('Features') # Changed 'features' to 'Features'
    if features and isinstance(features, list):
        doc.add_heading("Key Features", level=1) # Add a standard heading
        for feature in features:
            if isinstance(feature, str) and feature.strip():
                doc.add_paragraph(feature.strip(), style='List Bullet')

    # Table of Contents (array of objects) - Not in the provided user schema
    toc_items = article_dict.get('tableOfContents')
    if toc_items and isinstance(toc_items, list):
        doc.add_heading("Table of Contents", level=1)
        for item in toc_items:
            if isinstance(item, dict):
                text = item.get('text', '').strip()
                if text:
                    doc.add_paragraph(text, style='List Bullet')

    # Paragraphs (array of strings) - Not in the provided user schema as a top-level item
    paragraphs = article_dict.get('paragraphs') # This refers to a potential top-level 'paragraphs'
    if paragraphs and isinstance(paragraphs, list):
        for para in paragraphs:
            if isinstance(para, str) and para.strip():
                doc.add_paragraph(para.strip())

    # Steps (array of objects)
    steps_data = article_dict.get('Steps / How-To') # Changed 'steps' to 'Steps / How-To'
    if steps_data and isinstance(steps_data, list):
        doc.add_heading("Procedure / Steps", level=1)
        for i, step_item in enumerate(steps_data, 1):
            if isinstance(step_item, dict):
                step_text = step_item.get('step', '').strip()
                explanation = step_item.get('explanation', '').strip()
                screenshot_ref = step_item.get('screenshotRef', '').strip()

                if step_text:
                    p = doc.add_paragraph()
                    p.add_run(f"Step {i}: ").bold = True
                    p.add_run(step_text)

                    if explanation:
                         exp_p = doc.add_paragraph(explanation)
                         # Optional: Indent explanation
                         # exp_p.paragraph_format.left_indent = Inches(0.25)

                    if screenshot_ref:
                        ref_p = doc.add_paragraph()
                        ref_p.add_run(f"(See: {screenshot_ref})").italic = True
                        # Optional: Indent reference
                        # ref_p.paragraph_format.left_indent = Inches(0.25)

    # FAQ (array of objects)
    faq_data = article_dict.get('FAQ') # Changed 'faq' to 'FAQ'
    if faq_data and isinstance(faq_data, list):
        doc.add_heading("Frequently Asked Questions (FAQ)", level=1)
        for faq_item in faq_data:
            if isinstance(faq_item, dict):
                question = faq_item.get('question', '').strip()
                answer = faq_item.get('answer', '').strip()
                if question and answer:
                    p_q = doc.add_paragraph()
                    p_q.add_run("Q: ").bold = True
                    p_q.add_run(question)

                    p_a = doc.add_paragraph()
                    p_a.add_run("A: ").bold = True
                    p_a.add_run(answer)

    # Code Snippets (array of objects) - Not in the provided user schema
    code_snippets = article_dict.get('codeSnippets')
    if code_snippets and isinstance(code_snippets, list):
        doc.add_heading("Code Snippets", level=1)
        for snippet in code_snippets:
            if isinstance(snippet, dict):
                content = snippet.get('content', '').strip()
                language = snippet.get('language', 'plaintext').strip()
                caption = snippet.get('caption', '').strip()

                if content:
                    if caption:
                        doc.add_heading(caption, level=2)
                    lang_p = doc.add_paragraph()
                    lang_p.add_run(f"Language: {language}").italic = True
                    code_p = doc.add_paragraph(content, style='No Spacing')

    # Notes (array of strings) - Not in the provided user schema
    notes = article_dict.get('notes')
    if notes and isinstance(notes, list):
        doc.add_heading("Notes", level=1)
        for note in notes:
            if isinstance(note, str) and note.strip():
                doc.add_paragraph(note.strip(), style='Intense Quote')

    # Tips (array of strings) - Not in the provided user schema
    tips = article_dict.get('tips')
    if tips and isinstance(tips, list):
        doc.add_heading("Tips", level=1)
        for tip in tips:
            if isinstance(tip, str) and tip.strip():
                 doc.add_paragraph(tip.strip(), style='List Bullet')

    # Quotes (array of objects) - Not in the provided user schema
    quotes = article_dict.get('quotes')
    if quotes and isinstance(quotes, list):
        doc.add_heading("Quotes", level=1)
        for quote in quotes:
            if isinstance(quote, dict):
                text = quote.get('text', '').strip()
                attribution = quote.get('attribution', '').strip()
                if text:
                    doc.add_paragraph(text, style='Quote')
                    if attribution:
                        attr_p = doc.add_paragraph(f"— {attribution}")

    # Checklist (array of strings) - Not in the provided user schema
    checklist = article_dict.get('checklist')
    if checklist and isinstance(checklist, list):
        doc.add_heading("Checklist", level=1)
        for item in checklist:
            if isinstance(item, str) and item.strip():
                doc.add_paragraph(item.strip(), style='List Bullet')

    # Conclusion (object)
    conclusion = article_dict.get('Conclusion') # Changed 'conclusion' to 'Conclusion'
    if conclusion and isinstance(conclusion, dict):
        doc.add_heading("Conclusion", level=1)

        conc_paragraphs = conclusion.get('paragraphs')
        if conc_paragraphs and isinstance(conc_paragraphs, list):
            for para in conc_paragraphs:
                if isinstance(para, str) and para.strip():
                    doc.add_paragraph(para.strip())

        next_steps = conclusion.get('nextSteps')
        if next_steps and isinstance(next_steps, list):
            doc.add_heading("Next Steps", level=2)
            for step in next_steps:
                if isinstance(step, str) and step.strip():
                    doc.add_paragraph(step.strip(), style='List Bullet')

    # References (array of objects) - Not in the provided user schema
    references = article_dict.get('references')
    if references and isinstance(references, list):
        doc.add_heading("References", level=1)
        for ref in references:
            if isinstance(ref, dict):
                text = ref.get('text', '').strip()
                href = ref.get('href', '').strip()
                annotation = ref.get('annotation', '').strip()

                if text and href:
                    ref_text = f"{text}: {href}"
                    if annotation:
                        ref_text += f" ({annotation})"
                    doc.add_paragraph(ref_text, style='List Bullet')
                elif text:
                     doc.add_paragraph(text, style='List Bullet')

    # --- Save Document to Buffer ---
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer